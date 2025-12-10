#!/usr/bin/env python3
import argparse, json, os
from docx import Document
from docx.shared import Inches
ap=argparse.ArgumentParser(); ap.add_argument("--doc",required=True); ap.add_argument("--metrics",required=True); ap.add_argument("--pr",required=False); a=ap.parse_args()
doc=Document(a.doc)
doc.add_heading("7. Results (Updated)", level=1)
with open(a.metrics) as f: m=json.load(f)
doc.add_paragraph(f"mAP@0.5: {m.get('mAP50')}  |  mAP@0.5:0.95: {m.get('mAP50-95')}  |  Precision: {m.get('precision')}  |  Recall: {m.get('recall')}")
if a.pr and os.path.exists(a.pr): doc.add_paragraph('PR Curve:').add_run().add_picture(a.pr, width=Inches(5.5))
doc.save(a.doc); print("[✓] Updated", a.doc)
